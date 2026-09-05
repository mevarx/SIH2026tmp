"""
Task management and execution API routes.

Supports:
- POST /api/tasks (dispatch background task or synchronous processing)
- GET /api/tasks (list recent tasks)
- GET /api/tasks/{task_id} (fetch task status and output)
- DELETE /api/tasks/{task_id} (cancel running task)
- GET /api/tasks/{task_id}/stream (SSE token and reasoning stream)
"""

import asyncio
import json
import logging
import os
import re
import time
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, AsyncIterator, Dict, List, Optional

import aiofiles
from fastapi import APIRouter, BackgroundTasks, File, HTTPException, Query, UploadFile, status
from fastapi.responses import StreamingResponse

from app.agent.graph import agent_graph
from app.agent.router import task_router
from app.agent.state import AgentState
from app.config import settings
from app.models.base import ChatMessage, GenerationRequest
from app.models.registry import model_registry
from app.models.vision import VisionClient
from app.rag.retriever import hybrid_retriever
from app.sandbox.docker_runner import DockerSandboxRunner
from app.schemas.events import CompletionEvent, ErrorEvent, ReasoningEvent, StatusEvent, TokenEvent
from app.schemas.response import APIResponse
from app.schemas.tasks import (
    TaskCancelResponse,
    TaskListResponse,
    TaskPriority,
    TaskRequest,
    TaskResponse,
    TaskResult,
    TaskStatus,
    TaskStatusResponse,
    TaskType,
)
from app.security.audit import audit_logger
from app.vision.middleware import vision_middleware
from app.vision.ocr import ocr_document

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/tasks")

# In-memory fast cache and persistence file path
_tasks_cache: Dict[str, TaskResponse] = {}
_task_streams: Dict[str, asyncio.Queue] = {}
_task_store_path = Path(settings.task_store_path)


def _ensure_store():
    _task_store_path.parent.mkdir(parents=True, exist_ok=True)


async def _save_task_to_store(task: TaskResponse):
    _ensure_store()
    try:
        async with aiofiles.open(_task_store_path, mode="a", encoding="utf-8") as f:
            await f.write(task.model_dump_json() + "\n")
    except Exception as exc:
        logger.error("Failed to append task to store: %s", exc)


def _load_tasks_from_store():
    if not _task_store_path.is_file():
        return
    try:
        with open(_task_store_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    data = json.loads(line)
                    task = TaskResponse(**data)
                    _tasks_cache[task.task_id] = task
                except Exception:
                    continue
    except Exception as exc:
        logger.warning("Could not read existing tasks store: %s", exc)


_load_tasks_from_store()


async def _execute_task_pipeline(task_id: str, request: TaskRequest):
    """Background task execution pipeline."""
    task = _tasks_cache.get(task_id)
    if not task:
        return

    task.status = TaskStatus.RUNNING
    start_time = time.monotonic()
    queue = _task_streams.get(task_id)

    def emit_event(event):
        if queue:
            try:
                queue.put_nowait(event.to_sse())
            except Exception:
                pass

    emit_event(StatusEvent(status="running", message="Task execution started", task_id=task_id))
    pipeline_type = task_router.route(request)

    try:
        # Audit log the prompt
        p_type_str = pipeline_type.value if hasattr(pipeline_type, "value") else str(pipeline_type)
        await audit_logger.log_model_call(
            prompt=request.prompt,
            model=request.model or settings.active_model_name,
            task_type=p_type_str,
        )

        result_payload: Any = None
        artifacts: List[Dict[str, Any]] = []

        # Consolidate candidate attachments
        candidate_paths = list(request.file_paths)
        if request.attachment_path and request.attachment_path not in candidate_paths:
            candidate_paths.append(request.attachment_path)

        # Multimodal Vision Pipeline: Detect and optimize image attachments via Pillow
        multimodal_parts, non_image_paths = await vision_middleware.process_attachments(
            file_paths=candidate_paths,
            prompt=request.prompt,
        )

        effective_prompt = request.prompt

        # Process remaining document attachments (PDF, DOCX, TXT) and run local OCR
        if non_image_paths:
            emit_event(StatusEvent(status="reading_documents", message="Processing document attachments & running OCR...", task_id=task_id))
            doc_excerpts = []
            for fp in non_image_paths:
                p = Path(fp)
                if p.is_file():
                    try:
                        suffix = p.suffix.lower()
                        if suffix in (".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
                            d_text = await ocr_document(str(p))
                        elif suffix in (".txt", ".md", ".csv", ".json"):
                            d_text = p.read_text(encoding="utf-8", errors="replace")
                        elif suffix in (".docx", ".doc"):
                            from docx import Document
                            doc = Document(str(p))
                            d_text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
                        else:
                            d_text = ""
                        if d_text:
                            doc_excerpts.append(f"--- Document: {p.name} ---\n{d_text[:15000]}")
                    except Exception as e:
                        logger.warning("Could not read attached file %s: %s", fp, e)
            if doc_excerpts:
                effective_prompt = (
                    f"Attached Document(s) Content (OCR Extracted):\n" + "\n\n".join(doc_excerpts) +
                    f"\n\nUser Question/Instruction:\n{request.prompt}"
                )

        if pipeline_type == TaskType.AGENT:
            state = AgentState(
                task_id=task_id,
                prompt=effective_prompt,
                max_steps=6,
            )
            completed_state = await agent_graph.run(state)
            result_payload = completed_state.final_output
            artifacts = completed_state.artifacts

        elif pipeline_type == TaskType.RAG:
            emit_event(StatusEvent(status="retrieving", message="Searching local knowledge base...", task_id=task_id))
            retrieved = []
            try:
                retrieved = await hybrid_retriever.retrieve(request.prompt, top_k=5)
            except Exception as exc:
                logger.warning("RAG retrieval failed (%s); proceeding with internal reasoning fallback.", exc)
                emit_event(StatusEvent(
                    status="retrieval_warning",
                    message="Vector search unavailable; answering with internal reasoning.",
                    task_id=task_id,
                ))

            if retrieved:
                context_str = "\n\n".join(
                    [f"[Source: {c.source}, Score: {c.score}]\n{c.text}" for c in retrieved]
                )
                prompt_with_context = f"Context from Sovereign Knowledge Base:\n{context_str}\n\nUser Question:\n{request.prompt}"
                sys_prompt = request.system_prompt or "Answer questions strictly based on the provided context."
            else:
                prompt_with_context = request.prompt
                sys_prompt = request.system_prompt or "You are Sovereign AI Workbench. Answer the user question comprehensively."

            client = model_registry.get_client(role="reasoning", model_id=request.model)

            gen_req = GenerationRequest(
                prompt=prompt_with_context,
                system_prompt=sys_prompt,
                temperature=request.temperature or 0.2,
            )
            response = await client.chat(gen_req)
            result_payload = response.content
            artifacts = [
                {"type": "rag_sources", "sources": [c.model_dump() for c in retrieved]}
            ] if retrieved else []

        elif pipeline_type == TaskType.VISION:
            emit_event(StatusEvent(status="vision_processing", message="Vision Tower analyzing optimized multimodal image payload...", task_id=task_id))
            if multimodal_parts:
                client = model_registry.get_client(role="vision", model_id=request.model)
                gen_req = GenerationRequest(
                    model=request.model,
                    messages=[ChatMessage(role="user", content=multimodal_parts)],
                    temperature=request.temperature or 0.2,
                )
                resp = await client.chat(gen_req)
                result_payload = resp.content
            elif candidate_paths:
                vision_client = VisionClient()
                with open(candidate_paths[0], "rb") as img_f:
                    img_bytes = img_f.read()
                resp = await vision_client.analyze_image(
                    image_bytes=img_bytes,
                    prompt=request.prompt,
                    temperature=request.temperature or 0.2,
                )
                result_payload = resp.content
            else:
                result_payload = "Vision task requested but no image file attachments provided."

        elif pipeline_type == TaskType.SANDBOX:
            runner = DockerSandboxRunner()
            explicit_code = request.parameters.get("code")

            if not explicit_code:
                emit_event(StatusEvent(status="generating_code", message="Synthesizing executable Python code...", task_id=task_id))
                client = model_registry.get_client(role="coding", model_id=request.model)
                gen_prompt = (
                    f"Write a clean, self-contained Python 3 script that satisfies the following request:\n"
                    f"{request.prompt}\n\n"
                    f"Ensure you include print() calls demonstrating the output. Enclose your code in a ```python ... ``` block."
                )
                gen_resp = await client.chat(GenerationRequest(
                    prompt=gen_prompt,
                    temperature=0.2,
                ))
                llm_response_text = gen_resp.content or ""

                code_match = re.search(r"```(?:python)?\s*\n(.*?)\n```", llm_response_text, re.DOTALL)
                code_to_run = code_match.group(1).strip() if code_match else llm_response_text.strip()

                emit_event(StatusEvent(status="sandboxing", message="Executing in isolated sandbox...", task_id=task_id))
                sandbox_res = await runner.run_code(code_to_run, language="python")

                stdout_display = sandbox_res.stdout.strip() if sandbox_res.stdout.strip() else "(No standard output)"
                stderr_clean = "\n".join([line for line in sandbox_res.stderr.splitlines() if "not detected" not in line]).strip()
                stderr_display = f"\nStderr:\n{stderr_clean}" if stderr_clean else ""

                result_payload = (
                    f"{llm_response_text}\n\n"
                    f"### Execution Results:\n"
                    f"```text\n{stdout_display}{stderr_display}\n```\n"
                    f"*Exit code: {sandbox_res.exit_code} | Time: {sandbox_res.duration_seconds}s*"
                )
            else:
                code_to_run = explicit_code
                emit_event(StatusEvent(status="sandboxing", message="Executing code in isolated sandbox...", task_id=task_id))
                sandbox_res = await runner.run_code(code_to_run, language="python")
                result_payload = (
                    f"### Sandbox Execution Results:\n"
                    f"```text\n{sandbox_res.stdout or sandbox_res.stderr}\n```\n"
                    f"*Exit code: {sandbox_res.exit_code} | Time: {sandbox_res.duration_seconds}s*"
                )

            artifacts = [{
                "type": "sandbox_execution",
                "code": code_to_run,
                "stdout": sandbox_res.stdout,
                "stderr": sandbox_res.stderr,
                "exit_code": sandbox_res.exit_code,
            }]

        else:
            # GENERAL / LLM reasoning
            client = model_registry.get_client(role="reasoning", model_id=request.model)
            gen_req = GenerationRequest(
                prompt=effective_prompt,
                system_prompt=request.system_prompt,
                temperature=request.temperature or 0.7,
            )
            response = await client.chat(gen_req)
            result_payload = response.content
            if response.reasoning_content:
                emit_event(ReasoningEvent(reasoning=response.reasoning_content, task_id=task_id))

        duration = time.monotonic() - start_time
        task.status = TaskStatus.COMPLETED
        task.completed_at = datetime.now(timezone.utc)
        task.execution_time_seconds = round(duration, 3)
        task.result = TaskResult(
            output=result_payload,
            artifacts=artifacts,
            model_used=request.model or settings.active_model_name,
            execution_time_seconds=round(duration, 3),
        )

        emit_event(TokenEvent(token=str(result_payload), task_id=task_id))
        emit_event(CompletionEvent(result=result_payload, execution_time_seconds=round(duration, 3), task_id=task_id))

    except Exception as exc:
        logger.error("Task %s failed: %s", task_id, exc, exc_info=True)
        duration = time.monotonic() - start_time
        task.status = TaskStatus.FAILED
        task.completed_at = datetime.now(timezone.utc)
        task.execution_time_seconds = round(duration, 3)
        task.error = str(exc)
        emit_event(ErrorEvent(error=str(exc), task_id=task_id))

    finally:
        await _save_task_to_store(task)
        if queue:
            # Send end marker to signal SSE stream finish
            await queue.put(None)


@router.post("", response_model=TaskResponse, status_code=status.HTTP_202_ACCEPTED)
async def create_task(request: TaskRequest, background_tasks: BackgroundTasks) -> TaskResponse:
    """Dispatches a task for processing."""
    task = TaskResponse(
        task_type=request.task_type,
        status=TaskStatus.PENDING,
        prompt=request.prompt,
        metadata=request.metadata,
    )
    _tasks_cache[task.task_id] = task
    _task_streams[task.task_id] = asyncio.Queue()

    # Schedule non-blocking execution via BackgroundTasks
    background_tasks.add_task(_execute_task_pipeline, task.task_id, request)

    return task


@router.get("", response_model=TaskListResponse)
async def list_tasks(
    status_filter: Optional[TaskStatus] = Query(None, alias="status"),
    limit: int = Query(50, ge=1, le=200),
) -> TaskListResponse:
    """Lists all stored tasks with optional status filtering."""
    tasks = list(_tasks_cache.values())
    if status_filter:
        tasks = [t for t in tasks if t.status == status_filter]
    # Return newest first
    tasks.sort(key=lambda t: t.created_at, reverse=True)
    return TaskListResponse(tasks=tasks[:limit], total=len(tasks))


@router.post("/upload", response_model=APIResponse[Dict[str, Any]])
async def upload_task_document(file: UploadFile = File(...)) -> APIResponse[Dict[str, Any]]:
    """
    Direct in-chat document/image upload for local OCR & sovereign analysis.
    Saves to data/uploads/ and performs automatic text/OCR extraction without external network.
    """
    upload_dir = Path("data/uploads")
    upload_dir.mkdir(parents=True, exist_ok=True)

    filename = file.filename or "attachment"
    safe_name = re.sub(r"[^a-zA-Z0-9_.-]", "_", filename)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    target_path = upload_dir / f"{timestamp}_{safe_name}"

    try:
        with open(target_path, "wb") as f:
            shutil.copyfileobj(file.file, f)

        extracted_text = ""
        ocr_applied = False
        suffix = target_path.suffix.lower()

        if suffix in (".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
            extracted_text = await ocr_document(str(target_path))
            ocr_applied = True
        elif suffix in (".txt", ".md", ".csv", ".json"):
            extracted_text = target_path.read_text(encoding="utf-8", errors="replace")
        elif suffix in (".docx", ".doc"):
            from docx import Document
            doc = Document(str(target_path))
            extracted_text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())

        return APIResponse.ok(data={
            "filename": filename,
            "file_path": str(target_path).replace("\\", "/"),
            "file_size": target_path.stat().st_size,
            "extracted_text": extracted_text[:2000],
            "char_count": len(extracted_text),
            "ocr_applied": ocr_applied,
        })
    except Exception as exc:
        logger.error("Upload/OCR processing failed for %s: %s", filename, exc)
        return APIResponse.fail(error=f"Upload failed: {str(exc)}")


@router.get("/{task_id}", response_model=TaskResponse)
async def get_task(task_id: str) -> TaskResponse:
    """Fetches full task status and result."""
    task = _tasks_cache.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail=f"Task '{task_id}' not found")
    return task


@router.delete("/{task_id}", response_model=TaskCancelResponse)
async def cancel_task(task_id: str) -> TaskCancelResponse:
    """Cancels a pending or running task."""
    task = _tasks_cache.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail=f"Task '{task_id}' not found")

    if task.status in (TaskStatus.COMPLETED, TaskStatus.FAILED):
        return TaskCancelResponse(
            task_id=task_id,
            status=task.status,
            message="Task already finished; cannot cancel.",
        )

    task.status = TaskStatus.CANCELLED
    task.completed_at = datetime.now(timezone.utc)
    await _save_task_to_store(task)
    return TaskCancelResponse(task_id=task_id, status=TaskStatus.CANCELLED)


@router.get("/{task_id}/stream")
async def stream_task_events(task_id: str):
    """
    SSE streaming endpoint for real-time task progress and token generation.
    Connect via EventSource or fetch ReadableStream on client.
    """
    task = _tasks_cache.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail=f"Task '{task_id}' not found")

    queue = _task_streams.get(task_id)
    if not queue:
        queue = asyncio.Queue()
        _task_streams[task_id] = queue

    async def event_generator() -> AsyncIterator[str]:
        # If task already completed before stream connected, emit completion directly
        if task.status in (TaskStatus.COMPLETED, TaskStatus.FAILED, TaskStatus.CANCELLED):
            res_val = task.result.model_dump() if hasattr(task.result, "model_dump") else (task.result or task.error)
            yield CompletionEvent(result=res_val, task_id=task_id).to_sse()
            return

        while True:
            item = await queue.get()
            if item is None:
                break
            yield item

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )
